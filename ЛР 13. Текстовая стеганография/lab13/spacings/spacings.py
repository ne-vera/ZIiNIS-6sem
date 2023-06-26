from docx import Document
import re

def to_bits(s):
    bits = bin(int.from_bytes(s.encode(encoding='utf-8', errors='surrogatepass'), 'big'))[2:]
    return bits.zfill(8 * ((len(bits) + 7) // 8))

def from_bits(bits):
    n = int(bits, 2)
    return n.to_bytes((n.bit_length() + 7) // 8, 'big').decode(encoding='utf-8', errors='surrogatepass') or '\0'

def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def save_paragraphs_to_docx(paragraphs, filename):
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph.strip())
    document.save(filename)

def encode(message, container_path, saving_path):
    container_text = getText(container_path)
    
    bits = to_bits(message)
    space_count = container_text.count(' ')
    if space_count < int(len(bits) * 2):
        raise Exception('Not enough spaces in container!')
    
    # чередование одинарного пробела и двойного (хх_хх__хх) кодирует «1», 
    # переход же с двойного пробела на одинарный кодирует «0» (хх__хх_хх)
    spaces = []
    for bit in bits:
        if bit == '1':
            spaces.append(' ')
            spaces.append(' ' * 2)
        elif bit == '0':
            spaces.append(' ' * 2)
            spaces.append(' ')
    
    text_list = list(container_text)
    index = 0
    for i, char in enumerate(text_list):
        if char == ' ':
            space = spaces[index]
            text_list[i] = space
            index += 1
            if index >= len(spaces):
                break

    result = ''.join(text_list)    
    paragraphs = result.splitlines()
    save_paragraphs_to_docx(paragraphs, saving_path)
    print('Модифицированный текст сохранен в', saving_path)

# подсчет количества идущих подряд пробелов
def count_consecutive_spaces(string):
    pattern = r' +'
    matches = re.findall(pattern, string)
    counts = [len(match) for match in matches]
    return counts

def decode(container_path):
    container_text = getText(container_path)
    consecutive_spaces = count_consecutive_spaces(container_text)
    groups = [consecutive_spaces[i:i+2] for i in range(0, len(consecutive_spaces), 2)]
    bit_decoded = ''
    for group in groups:
        if group == [2,1]:
            bit_decoded += '0'
        elif group == [1,2]:
            bit_decoded += '1'
    decoded = from_bits(bit_decoded)
    print('Расшифрованное сообщение:', decoded)
